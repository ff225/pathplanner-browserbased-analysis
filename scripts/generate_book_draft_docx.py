#!/usr/bin/env python3
"""Generate Bab 5 Use Case 6 Smart Air Quality draft as .docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_fill="D9E2F3"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def main():
    out = Path(__file__).resolve().parents[1] / "docs" / "Bab5_UseCase6_Smart_Air_Quality_Bu_Rini.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Bab 5 (cuplikan) — Use Case 6: Smart Air Quality", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Kontributor: Bu Rini Apriyanti Purba")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph(
        "Draf pertama untuk buku «Sensing for Smart X». "
        "Berdasarkan riset PathPlanner (Crowdsensing for Health and Environment)."
    )

    # 6.1
    doc.add_heading("6.1 Konteks: dari sensor udara ke keputusan mobilitas sehat", level=1)
    doc.add_paragraph(
        "Smart Environment tidak berhenti pada monitoring kualitas udara (KU). Nilai tambah muncul "
        "ketika data sensing—dari stasiun resmi, layanan terbuka, peta partisipatif, hingga (ke depan) "
        "kontribusi warga—dijadikan input keputusan yang dapat dioperasikan: peringatan, visualisasi, "
        "dan rute pejalan kaki yang meminimalkan paparan tanpa mengabaikan waktu tempuh."
    )
    doc.add_paragraph(
        "Proyek PathPlanner (Crowdsensing for Health and Environment) mengilustrasikan pola umum buku ini:"
    )
    p = doc.add_paragraph()
    p.add_run("Akuisisi data (alat & non-alat) → pra-pemrosesan → layanan data → aplikasi (peta, skor, rute) → monitoring & analisis.").bold = True
    doc.add_paragraph(
        "Berbeda dari use case yang bergantung pada satu jenis sensor fisik di lokasi tetap, "
        "Smart Air Quality di kota besar biasanya hibrida, seperti di Tabel 6.1."
    )
    doc.add_paragraph("Tabel 6.1 — Sumber data dalam PathPlanner", style="Intense Quote")
    add_table(
        doc,
        ["Sumber", "Peran sensing", "Contoh di PathPlanner"],
        [
            ["Stasiun / API kualitas udara", "PM₂.₅, indeks polusi", "OpenAQ, WHO"],
            ["Meteorologi", "suhu, kelembaban", "Open-Meteo"],
            ["Topografi / jalan", "kemiringan (slope)", "OpenTopoData, Mapbox"],
            ["Peta terbuka", "kebisingan perkiraan, ruang hijau", "OpenStreetMap (Overpass)"],
            ["Routing", "geometri jalan, jarak, waktu", "Mapbox"],
            ["Profil pengguna", "prioritas faktor lingkungan", "6 profil kesehatan"],
        ],
    )
    doc.add_paragraph(
        "Inti Bab 1–2 tetap berlaku: sensor (dan sumber data setara sensor) bukan tujuan akhir, "
        "melainkan bahan baku sistem cerdas."
    )

    # 6.2
    doc.add_heading("6.2 Sensor / Akuisisi Data", level=1)
    doc.add_heading("6.2.1 Data lingkungan multi-faktor", level=2)
    doc.add_paragraph(
        "PathPlanner mengukur paparan sepanjang rute, bukan hanya di satu titik. Untuk setiap pasangan "
        "asal–tujuan, sistem: (1) menghasilkan geometri rute via Mapbox; (2) mensampling 6–24 titik "
        "sepanjang polyline (mode benchmark: 8 titik); (3) memanggil API lingkungan Django yang "
        "mengagregasi kualitas udara, suhu, kelembaban, kebisingan, dan kemiringan."
    )
    doc.add_paragraph(
        "Pendekatan ini mendekati mobile crowdsensing tanpa mengharuskan setiap pejalan membawa sensor: "
        "telepon atau browser sebagai platform akuisisi kontekstual di sepanjang lintasan mobilitas."
    )
    doc.add_heading("6.2.2 Data non-alat (social & open data)", level=2)
    doc.add_paragraph(
        "OpenStreetMap menyediakan morfologi jalan dan proxy lingkungan. Profil kondisi kesehatan "
        "menggunakan bobot bukti (evidence tiers 1/2/3). Benchmark terstruktur: 360 evaluasi rute "
        "(6 kota × 6 profil × 10 pasangan OD pejalan kaki, jarak langsung 1–3 km)."
    )
    doc.add_paragraph(
        "Kota percontohan: Tokyo, Shanghai, New York, London, Barcelona, Jakarta."
    )
    doc.add_heading("6.2.3 Arah crowdsensing partisipatif", level=2)
    doc.add_paragraph(
        "PathPlanner/Crowdsensing dirancang agar data warga (ke depan) melengkapi stasiun resmi yang "
        "jarang dan tidak seragam spasial. Implementasi saat ini mengandalkan sensor jaringan dan open data, "
        "dengan validasi sistematis via benchmark browser (Playwright)."
    )

    # 6.3
    doc.add_heading("6.3 Fungsi (peran dalam Smart X)", level=1)
    doc.add_heading("6.3.1 Monitoring & visualisasi", level=2)
    doc.add_paragraph(
        "Dashboard peta menampilkan layer kualitas udara, cuaca, kebisingan, dan kemiringan. "
        "Skor lingkungan rute Uc(π) pada skala 1–10 (lebih tinggi = lebih sehat untuk kondisi c):"
    )
    doc.add_paragraph("Uc(π) = Σ w_c,k · f_k(x̄_π,k) / Σ w_c,k", style="Intense Quote")
    doc.add_heading("6.3.2 Perencanaan rute berbasis paparan", level=2)
    doc.add_paragraph(
        "Alur fungsional (selaras Bab 3): browser/UI → generator kandidat (A* lingkungan + waypoint) → "
        "Mapbox routing → scoring lima faktor → kebijakan detour ≤ 5 km (+ toleransi lunak 0,5 km). "
        "Waktu tempuh tidak dioptimalkan langsung pada tahap seleksi."
    )
    doc.add_heading("6.3.3 Analisis & pengambilan keputusan", level=2)
    doc.add_paragraph("Tabel 6.2 — Tiga lensa analisis riset", style="Intense Quote")
    add_table(
        doc,
        ["Analisis", "Pertanyaan", "Metrik"],
        [
            ["1. Headroom antar-kota", "Di kota mana sistem punya ruang perbaikan?", "Mean detour, % strict gain, % fallback"],
            ["2. Trade-off paparan–waktu", "Kapan rekomendasi layak deploy?", "EI, TP, Eff = EI − TP, kuadran"],
            ["3. Robustness profil", "Performa stabil antar kondisi?", "Condition Sensitivity Index (CSI)"],
        ],
    )

    # 6.4
    doc.add_heading("6.4 Manfaat", level=1)
    doc.add_heading("6.4.1 Manfaat langsung", level=2)
    benefits = [
        "Kesadaran paparan mikro: dua rute dengan jarak mirip dapat berbeda secara udara, panas, kebisingan, dan kemiringan.",
        "Personalisasi kesehatan lingkungan: rute optimal berbeda per profil (respiratory vs mental, dll.).",
        "Kebijakan kota: polusi rata-rata ≠ manfaat routing yang terealisasi.",
        "Reproduksibilitas ilmiah: benchmark Playwright menilai stack yang benar-benar di-deploy.",
    ]
    for b in benefits:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("6.4.2 Temuan empiris (benchmark browser, 360 rute)", level=2)
    doc.add_paragraph("Tabel 6.3 — Ringkasan global", style="Intense Quote")
    add_table(
        doc,
        ["Indikator", "Nilai"],
        [
            ["Peningkatan skor lingkungan rata-rata (EI)", "0,53%"],
            ["Penalti waktu rata-rata (TP)", "28,0%"],
            ["Efisiensi rata-rata (Eff)", "−27,5"],
            ["Rute dengan gain lingkungan ketat", "9,2%"],
            ["Rute dengan Eff > 0 (layak deploy)", "1,9%"],
            ["Fallback «langsung, tanpa gain»", "30%"],
        ],
    )
    doc.add_paragraph("Tabel 6.4 — Performa per kota", style="Intense Quote")
    add_table(
        doc,
        ["Kota", "Mean detour (m)", "Mean Eff", "% win", "CSI (↓ = robust)"],
        [
            ["London", "275", "−11,4", "5,0%", "5,7"],
            ["New York", "234", "−11,3", "3,3%", "16,8"],
            ["Tokyo", "247", "−11,2", "1,7%", "11,0"],
            ["Barcelona", "301", "−12,1", "0%", "5,3"],
            ["Shanghai", "449", "−21,9", "1,7%", "7,7"],
            ["Jakarta", "2058", "−96,9", "0%", "23,5"],
        ],
    )
    doc.add_paragraph(
        "Pesan pedagogis: Smart Air Quality routing bukan solusi ajaib; manfaat kecil pada perjalanan "
        "pendek 1–3 km. Jakarta menjadi stres uji jaringan dan routability, bukan sekadar «kota terpolusi»."
    )

    # 6.5
    doc.add_heading("6.5 Tambahan", level=1)
    doc.add_heading("6.5.1 Keterkaitan dengan Bab 1–4", level=2)
    add_table(
        doc,
        ["Bab", "Kaitan Smart Air Quality"],
        [
            ["1 Dasar sensing", "Multi-faktor; akurasi vs proxy OSM; delay API"],
            ["2 Peran dalam Smart X", "Input mobilitas sehat & keadilan paparan"],
            ["3 Arsitektur", "Device (browser) – komunikasi – data – aplikasi"],
            ["4 Manajemen data", "Time-series; cleaning saat API gagal; CSV benchmark"],
        ],
    )
    doc.add_heading("6.5.2 Tantangan", level=2)
    challenges = [
        "Kesenjangan tujuan: sistem memaksimalkan skor lingkungan; evaluasi memakai Eff = EI − TP.",
        "Cakupan pendek: 1–3 km pejalan kaki.",
        "Ketergantungan layanan pihak ketiga (OpenAQ, Mapbox, OSM).",
        "Keadilan antar profil: CSI Jakarta 23,5 vs Barcelona 5,3.",
        "Privasi: profil kesehatan + jejak lokasi memerlukan governance.",
    ]
    for c in challenges:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("6.5.3 Arah pengembangan (selaras Bab 6)", level=2)
    future = [
        "AI: pembelajaran bobot dari data paparan (dengan privasi).",
        "Digital Twin: simulasi skenario polusi sebelum deploy massal.",
        "Crowdsensing aktif: kalibrasi PM₂.₅ hyperlokal + fusi stasiun resmi.",
        "Optimasi multi-objektif: gabungkan Eff atau ambang gain minimum ke seleksi rute.",
    ]
    for f in future:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("6.5.4 Kotak refleksi untuk mahasiswa", level=2)
    doc.add_paragraph(
        "«Mengapa London punya PM₂.₅ lebih rendah dari Jakarta, tetapi CSI-nya tidak jauh lebih baik "
        "dari Shanghai? Apa peran morfologi jalan vs polusi absolut?»",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "«Jika 30% rute kembali ke Fallback Direct, apakah sistem gagal—orang bijak menolak detour tanpa gain?»",
        style="Intense Quote",
    )

    # 6.6
    doc.add_heading("6.6 Ringkasan satu halaman", level=1)
    doc.add_paragraph(
        "Smart Air Quality (PathPlanner) menggabungkan sensing lingkungan multi-sumber dan profil kesehatan "
        "untuk memantau, menilai, dan merekomendasikan rute pejalan kaki di enam metropolis global. "
        "Arsitekturnya mengikuti pola Smart X: data bersih dan terstruktur → skor transparan → keputusan "
        "rute dengan batas detour. Manfaat utamanya adalah literasi paparan, personalisasi, dan metode "
        "evaluasi reproduksibel; benchmark menunjukkan gain lingkungan rata-rata masih kecil dibanding "
        "penalti waktu, sehingga use case ini mengajarkan bahwa sensing canggih harus diikuti desain "
        "kebijakan dan metrik deployability."
    )

    doc.add_page_break()
    doc.add_heading("Catatan editorial untuk tim buku", level=1)
    notes = [
        "Panjang: ~2.500 kata; bisa dipendekkan ke 8–12 halaman dengan tabel di lampiran.",
        "Gambar disarankan: arsitektur 4 lapisan; peta contoh rute; diagram kuadran deployability.",
        "Bahasa: outline Indonesia; abstrak paper Inggris—opsi box Technical Note 1 halaman.",
        "Koherensi Bab 5: tekankan pola sensor → preprocessing → cloud → dashboard.",
    ]
    for n in notes:
        doc.add_paragraph(n, style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Sumber data: paper_overleaf.tex, analysis_outputs/, pathplanner-main-liam. ").italic = True
    footer.add_run(f"File: {out.name}").italic = True

    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
